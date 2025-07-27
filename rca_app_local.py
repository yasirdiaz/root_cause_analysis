import streamlit as st
from docx import Document
from docx.shared import Inches, RGBColor
from PIL import Image, ImageDraw
import datetime
import os
import io
import gspread
from google.oauth2.service_account import Credentials
import pandas as pd

# --- CONFIGURATION FOR GOOGLE SHEETS ---
scopes = [
    'https://www.googleapis.com/auth/spreadsheets.readonly'
]
try:
    credentials = Credentials.from_service_account_file('credentials.json', scopes=scopes)
    gspread_client = gspread.authorize(credentials)
except Exception as e:
    st.error(f"Error loading Google Sheets credentials: {e}")
    st.stop()

# RCA Health Report Sheet Configuration
GOOGLE_SHEET_RCA_REPORT_ID = "1-_rjoj_9bXUaa_ZZydEXs6wbUFQ_kjC9rYwAiDc-zXA"
GOOGLE_SHEET_RCA_REPORT_NAME = "Health Report"
GOOGLE_SHEET_REPORT_RANGE = "C8:J32"

QUEUE_COL_INDEX_IN_RANGE = 0
STATUS_COL_INDEX_IN_RANGE = 7

# Volume Data Google Sheet Configuration
VOLUME_SHEET_ID = "1MUcv83VOBUoEJQhIpsp5HXAtNaaT7-AJQEjlkL1_egs"
VOLUME_SHEET_NAME = "Sheet1"
VOLUME_QUEUE_COL_INDEX = 3
VOLUME_VALUE_COL_INDEX = 2


# --- HELPER FUNCTIONS ---

def create_colored_circle_image(color_name, output_dir):
    """Creates a small PNG image of a colored circle."""
    size = (24, 24)
    image = Image.new("RGBA", size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)

    colors_rgb = {
        "red": (255, 0, 0),
        "amber": (255, 191, 0),
        "green": (0, 128, 0)
    }

    rgb_color = colors_rgb.get(color_name.lower(), (0, 0, 0))

    radius = min(size) // 2 - 1
    center = (size[0] // 2, size[1] // 2)

    draw.ellipse((center[0] - radius, center[0] - radius,
                  center[0] + radius, center[1] + radius),
                 fill=rgb_color)

    os.makedirs(output_dir, exist_ok=True)

    image_path = os.path.join(output_dir, f"{color_name}_circle.png")
    image.save(image_path)
    return image_path


def map_sheet_color_char_to_name(sheet_status_char):
    """Maps Unicode Health Indicator characters (or text) to color names."""
    status_map = {
        "🔴": "red",
        "🟡": "amber",
        "🟢": "green",
        "red": "red",
        "amber": "amber",
        "green": "green"
    }
    return status_map.get(sheet_status_char.lower().strip(), None)


def get_all_relevant_queues_from_sheet(gspread_client, sheet_id, sheet_name, sheet_range, queue_col_idx_in_range,
                                       status_col_idx_in_range):
    """Reads the Google Sheet and filters for 'Red' or 'Amber' queues."""
    try:
        workbook = gspread_client.open_by_key(sheet_id)
        worksheet = workbook.worksheet(sheet_name)
        all_data_rows = worksheet.get(sheet_range)

        if not all_data_rows:
            st.warning(f"Error: The Google Sheet is empty or no data found in range '{sheet_range}'.")
            return []

        relevant_queues_list = []

        for r_idx, row in enumerate(all_data_rows):
            if len(row) <= status_col_idx_in_range:
                continue

            queue_name = row[queue_col_idx_in_range].strip()
            status_char = row[status_col_idx_in_range].strip()

            if not queue_name or queue_name.lower() in ["email queues", "live queues"]:
                continue

            mapped_color = map_sheet_color_char_to_name(status_char)

            if mapped_color and mapped_color in ["red", "amber"]:
                relevant_queues_list.append(
                    {"Queue Name": queue_name, "Status Color": mapped_color, "Range Index": r_idx})

        return relevant_queues_list

    except gspread.exceptions.SpreadsheetNotFound:
        st.error(f"Error: Google Sheet with ID '{sheet_id}' not found. Check ID and service account permissions.")
        return []
    except gspread.exceptions.WorksheetNotFound:
        st.error(f"Error: Worksheet with name '{sheet_name}' not found. Check sheet name.")
        return []
    except Exception as e:
        st.error(f"An unexpected error occurred while reading Google Sheet: {e}")
        return []


def get_actual_volume_from_sheet(gspread_client, sheet_id, sheet_name, queue_name_to_lookup, queue_col_idx,
                                 volume_col_idx):
    """Looks up the actual volume for a given queue name in a specified Google Sheet."""
    try:
        workbook = gspread_client.open_by_key(sheet_id)
        worksheet = workbook.worksheet(sheet_name)
        all_data_rows = worksheet.get_all_values()

        if not all_data_rows:
            st.warning(f"Error: Volume sheet '{sheet_name}' is empty or no data found.")
            return None

        for row in all_data_rows[1:]:
            if len(row) > volume_col_idx:
                sheet_queue_name = row[queue_col_idx].strip()
                sheet_volume_value = row[volume_col_idx].strip()

                if sheet_queue_name.lower() == queue_name_to_lookup.lower():
                    st.success(f"Volume found for '{queue_name_to_lookup}': {sheet_volume_value}")
                    return sheet_volume_value

        st.warning(f"Warning: Volume for queue '{queue_name_to_lookup}' not found in volume sheet.")
        return None

    except gspread.exceptions.SpreadsheetNotFound:
        st.error(f"Error: Volume Sheet with ID '{sheet_id}' not found. Check ID and service account permissions.")
        return None
    except gspread.exceptions.WorksheetNotFound:
        st.error(f"Error: Volume Worksheet with name '{sheet_name}' not found. Check sheet name.")
        return None
    except Exception as e:
        st.error(f"An unexpected error occurred while reading volume sheet: {e}")
        return None


# --- STREAMLIT APP LOGIC ---

st.set_page_config(layout="wide")
st.title("RCA Report Generator")

# Initialize session state variables
if 'rca_reports' not in st.session_state:
    st.session_state.rca_reports = []
if 'document' not in st.session_state:
    st.session_state.document = Document()
if 'remaining_queues' not in st.session_state:
    st.session_state.remaining_queues = None
if 'current_rca_step' not in st.session_state:
    st.session_state.current_rca_step = "select_queue"
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'driver_text_area_key' not in st.session_state:
    st.session_state.driver_text_area_key = 0
if 'temp_dir' not in st.session_state:
    st.session_state.temp_dir = "temp_report_images"
    os.makedirs(st.session_state.temp_dir, exist_ok=True)


# Function to process an RCA and add to document
def process_single_rca_to_document(rca_data):
    doc = st.session_state.document
    temp_dir = st.session_state.temp_dir

    try:
        if len(st.session_state.rca_reports) > 0:
            doc.add_paragraph()
            doc.add_paragraph()

            # Add Content to Word Document
        circle_image_path = create_colored_circle_image(rca_data['report_color'], temp_dir)

        # Title
        p_title = doc.add_paragraph()
        run_img = p_title.add_run()
        run_img.add_picture(circle_image_path, width=Inches(0.18), height=Inches(0.18))
        run_text_title = p_title.add_run(f" {rca_data['title_skill'].upper()}")
        run_text_title.font.bold = True
        run_text_title.font.color.rgb = RGBColor(0, 0, 0)
        p_title.style = 'Heading 1'

        # Supply Section
        p_supply_subtitle = doc.add_paragraph()
        run_text_supply_subtitle = p_supply_subtitle.add_run("Supply")
        run_text_supply_subtitle.font.bold = True
        run_text_supply_subtitle.font.color.rgb = RGBColor(0, 0, 0)
        p_supply_subtitle.style = 'Heading 2'
        for line_text in rca_data['supply_word_lines']:
            p = doc.add_paragraph(line_text)
            p.style = 'List Bullet'

        # Demand Section
        p_demand_subtitle = doc.add_paragraph()
        run_text_demand_subtitle = p_demand_subtitle.add_run("Demand")
        run_text_demand_subtitle.font.bold = True
        run_text_demand_subtitle.font.color.rgb = RGBColor(0, 0, 0)
        p_demand_subtitle.style = 'Heading 2'
        doc.add_paragraph(f"Forecast volume was {rca_data['forecasted_volume_str']}.").style = 'List Bullet'
        doc.add_paragraph(f"Actual volume was {rca_data['actual_volume_str']}.").style = 'List Bullet'
        doc.add_paragraph(rca_data['variance_text']).style = 'List Bullet'

        # Main Drivers & Mitigation actions Section
        p_drivers_subtitle = doc.add_paragraph()
        run_text_drivers_subtitle = p_drivers_subtitle.add_run("Main Drivers & Mitigation actions")
        run_text_drivers_subtitle.font.bold = True
        run_text_drivers_subtitle.font.color.rgb = RGBColor(0, 0, 0)
        p_drivers_subtitle.style = 'Heading 2'
        for driver_item in rca_data['drivers_list']:
            p_driver = doc.add_paragraph(driver_item)
            p_driver.style = 'List Bullet'

        st.session_state.rca_reports.append(rca_data)
        st.session_state.messages.append(f"RCA for '{rca_data['title_skill']}' added to document.")

    except Exception as e:
        st.error(f"Error adding RCA to document: {e}")
    finally:
        if circle_image_path and os.path.exists(circle_image_path):
            os.remove(circle_image_path)
            try:
                os.rmdir(temp_dir)
            except OSError:
                pass


# --- Streamlit UI Flow ---

if st.session_state.current_rca_step == "select_queue":
    st.header("Step 1: Select Queue for RCA")
    if st.session_state.remaining_queues is None:
        st.session_state.remaining_queues = get_all_relevant_queues_from_sheet(
            gspread_client,
            GOOGLE_SHEET_RCA_REPORT_ID,
            GOOGLE_SHEET_RCA_REPORT_NAME,
            GOOGLE_SHEET_REPORT_RANGE,
            QUEUE_COL_INDEX_IN_RANGE,
            STATUS_COL_INDEX_IN_RANGE
        )

    # Debugging output - remove after confirming fix
    # st.write("--- Debugging get_all_relevant_queues_from_sheet ---")
    # st.write(f"Result from get_all_relevant_queues_from_sheet: {st.session_state.remaining_queues}")
    # st.write("--- End Debugging ---")

    if not st.session_state.remaining_queues:
        st.warning(
            "No Red or Amber queues found to process. Please update your Google Sheet or check sheet configuration.")
        if len(st.session_state.rca_reports) > 0:
            st.subheader("All available RCAs processed or none found.")
            st.session_state.current_rca_step = "finish"
        else:
            st.info("No RCAs could be started. No document will be generated.")
            st.stop()
    else:
        st.write("--- Remaining Queues for RCA (Red or Amber status) ---")
        # --- ADDED CODE TO DISPLAY THE LIST ---
        for i, item in enumerate(st.session_state.remaining_queues):
            st.write(f"{i + 1}. {item['Queue Name']}, color {item['Status Color'].capitalize()}")
        # --- END ADDED CODE ---

        selected_option_str = st.text_input(
            "Which queue number would you like to process now? (Type 'Finish' to exit and save)",
            key="queue_selector_text_input"
        ).strip().lower()

        if selected_option_str:
            if selected_option_str == "finish":
                st.session_state.current_rca_step = "finish"
                st.rerun()

            try:
                choice_idx = int(selected_option_str) - 1
                if 0 <= choice_idx < len(st.session_state.remaining_queues):
                    selected_item = st.session_state.remaining_queues[choice_idx]

                    st.session_state.rca_data = {
                        'title_skill': selected_item['Queue Name'],
                        'report_color': selected_item['Status Color'],
                        'selected_row_index_in_range': selected_item['Range Index']
                    }
                    st.session_state.remaining_queues.pop(choice_idx)

                    st.session_state.current_rca_step = "collect_inputs"
                    st.rerun()
                else:
                    st.error("Invalid number. Please choose from the list.")
            except ValueError:
                st.error("Invalid input. Please enter a number or type 'Finish'.")


elif st.session_state.current_rca_step == "collect_inputs":
    st.header(f"Step 2: Collect RCA Details for '{st.session_state.rca_data['title_skill']}'")

    selected_row_index = st.session_state.rca_data['selected_row_index_in_range']
    if 1 <= selected_row_index <= 13:
        rca_type = "emails"
    elif 17 <= selected_row_index <= 24:
        rca_type = "phone"
    else:
        st.warning(
            f"Queue '{st.session_state.rca_data['title_skill']}' is from an unexpected row index (range index: {selected_row_index}). Defaulting to 'phone/chats' logic. Please verify the sheet range mapping.")
        rca_type = "phone"
    st.info(f"RCA Type auto-detected: **{rca_type.capitalize()}**")
    st.session_state.rca_data['rca_type'] = rca_type

    st.subheader("Supply Section")
    supply_word_lines = []

    if rca_type == "emails":
        required_email_hours = st.text_input(
            "How many required email hours to handle incoming volumes and existing backlog?", key="req_email_hours")
        required_email_hours_new_demand = st.text_input("How many required email hours to handle incoming volumes?",
                                                        key="req_email_new_demand")
        actual_email_hours = st.text_input("How many supply hours?", key="actual_email_hours")

        if required_email_hours and required_email_hours_new_demand and actual_email_hours:
            supply_word_lines.append(
                f"Required email hours to handle incoming volumes and existing backlog was {required_email_hours}.")
            supply_word_lines.append(
                f"Required email hours to handle incoming volumes was {required_email_hours_new_demand}.")
            supply_word_lines.append(f"Actual email hours were {actual_email_hours}.")
            st.session_state.rca_data['supply_word_lines'] = supply_word_lines
        else:
            st.session_state.rca_data['supply_word_lines'] = []
            st.warning("Please fill all Supply (Email) fields.")

    else:
        required_hours_live_channels_str = st.text_input("How many required hours to handle forecasted volume?",
                                                         key="req_live_hours")
        actual_live_channel_hours_str = st.text_input("How many actual hours were?", key="actual_live_hours")

        result_text = ""
        if required_hours_live_channels_str and actual_live_channel_hours_str:
            try:
                required_hours_live_channels = float(required_hours_live_channels_str)
                actual_live_channel_hours = float(actual_live_channel_hours_str)

                if required_hours_live_channels == 0:
                    result_text = "Cannot calculate variance: Required hours to handle forecasted volume is zero."
                else:
                    result_value = (actual_live_channel_hours / required_hours_live_channels) - 1
                    result_formatted = f"{abs(result_value):.2%}"

                    if actual_live_channel_hours > required_hours_live_channels:
                        result_text = f"Incoming volume was {result_formatted} higher than forecasted."
                    else:
                        result_text = f"Incoming volume was {result_formatted} lower than forecasted."
            except ValueError:
                result_text = "Error: Invalid numeric input. Cannot calculate variance."

            supply_word_lines.append(
                f"Required hours to handle forecasted volume were {required_hours_live_channels_str}.")
            supply_word_lines.append(f"Actual hours were {actual_live_channel_hours_str}.")
            supply_word_lines.append(result_text)
            st.session_state.rca_data['supply_word_lines'] = supply_word_lines
        else:
            st.session_state.rca_data['supply_word_lines'] = []
            st.warning("Please fill all Supply (Live Channels) fields.")

    st.subheader("Demand Section")
    forecasted_volume_str = st.text_input("What is the forecasted volume?", key="forecasted_volume")

    actual_volume_str_auto = get_actual_volume_from_sheet(
        gspread_client,
        VOLUME_SHEET_ID,
        VOLUME_SHEET_NAME,
        st.session_state.rca_data['title_skill'],
        VOLUME_QUEUE_COL_INDEX,
        VOLUME_VALUE_COL_INDEX
    )
    actual_volume_str_manual = st.text_input(
        f"Actual volume (auto-fetched for '{st.session_state.rca_data['title_skill']}'):",
        value=actual_volume_str_auto if actual_volume_str_auto is not None else "",
        key="actual_volume_input"
    )
    actual_volume_str = actual_volume_str_manual

    variance_text = ""
    if forecasted_volume_str and actual_volume_str:
        try:
            forecasted_volume = float(forecasted_volume_str)
            actual_volume = float(actual_volume_str)

            if forecasted_volume == 0:
                variance_text = "Cannot calculate variance: Forecasted volume is zero."
            else:
                variance_value = (actual_volume / forecasted_volume) - 1
                variance_formatted = f"{abs(variance_value):.2%}"

                if actual_volume > forecasted_volume:
                    variance_text = f"Incoming volume was {variance_formatted} higher than forecasted."
                else:
                    variance_text = f"Incoming volume was {variance_formatted} lower than forecasted."
        except ValueError:
            variance_text = "Error: Invalid numeric input for forecasted or actual volume. Cannot calculate variance."
        st.session_state.rca_data['variance_text'] = variance_text
    else:
        st.session_state.rca_data['variance_text'] = ""
        st.warning("Please fill all Demand fields.")

    st.subheader("Main Drivers & Mitigation actions")

    if 'drivers_list' not in st.session_state.rca_data:
        st.session_state.rca_data['drivers_list'] = []

    st.write("What were the main drivers and/or mitigation actions?")
    driver_input = st.text_area("Type an action/driver and click 'Add Action'",
                                key=f"driver_text_area_{st.session_state.driver_text_area_key}")

    if st.button("Add Action", key="add_driver_button"):
        if driver_input.strip():
            st.session_state.rca_data['drivers_list'].append(driver_input.strip())
            st.session_state.messages.append(f"Added driver: {driver_input.strip()}")
            st.session_state.driver_text_area_key += 1
            st.rerun()

    if st.session_state.rca_data['drivers_list']:
        st.write("--- Current Drivers/Actions ---")
        for i, driver in enumerate(st.session_state.rca_data['drivers_list']):
            st.write(f"- {driver}")
    else:
        st.info("No drivers/actions added yet.")

    st.session_state.rca_data['forecasted_volume_str'] = forecasted_volume_str
    st.session_state.rca_data['actual_volume_str'] = actual_volume_str

    if st.button("Complete RCA & Add to Report", key="complete_rca_button"):
        if st.session_state.rca_data['supply_word_lines'] and \
                st.session_state.rca_data['variance_text'] and \
                st.session_state.rca_data['drivers_list']:
            process_single_rca_to_document(st.session_state.rca_data)
            st.session_state.current_rca_step = "review_rca"
            st.rerun()
        else:
            st.error("Please ensure all sections (Supply, Demand, Main Drivers) are filled for the current RCA.")


elif st.session_state.current_rca_step == "review_rca":
    st.header("Step 3: Review & Finalize Report")
    st.success(f"RCA for '{st.session_state.rca_data['title_skill']}' has been added to the report!")

    if len(st.session_state.remaining_queues) > 0:
        st.write(f"You have {len(st.session_state.remaining_queues)} queues remaining to process.")
        if st.button("Process Next RCA", key="next_rca_button"):
            st.session_state.current_rca_step = "select_queue"
            st.rerun()
    else:
        st.info("All Red or Amber queues from the sheet have been processed!")

    if st.button("Finalize & Download Report", key="finalize_download_button"):
        st.session_state.current_rca_step = "finish"
        st.rerun()

    if st.button("Start Over", key="start_over_button"):
        st.session_state.clear()
        st.rerun()


elif st.session_state.current_rca_step == "finish":
    st.header("RCA Processing Complete")
    st.info("No more Red or Amber queues found to process.")

    if len(st.session_state.rca_reports) > 0:
        output_filename = f"RCA_Multi_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        bio = io.BytesIO()
        st.session_state.document.save(bio)
        bio.seek(0)
        st.download_button(
            label="Download Final Report",
            data=bio,
            file_name=output_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_finished_report"
        )
        st.success("Your report is ready for download!")
    else:
        st.warning("No RCAs were processed. No report to download.")

    if st.button("Start New Session", key="new_session_button"):
        st.session_state.clear()
        st.rerun()

# Display Messages (e.g., recent activity/status)
if st.session_state.messages:
    st.sidebar.subheader("Recent Activity")
    for msg in reversed(st.session_state.messages):
        st.sidebar.markdown(f"- {msg}")
